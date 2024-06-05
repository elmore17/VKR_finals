import json
from docx import Document
from docxtpl import DocxTemplate
import locale
from datetime import datetime
import psycopg2

# Устанавливаем русскую локаль
locale.setlocale(locale.LC_TIME, 'Russian_Russia.1251')

def get_db_connection():
    return psycopg2.connect(database="kisprod", user="postgres", password="elmore", host="localhost", port="5432")

def read_docx(docx_path, items_per_subarray=142):
    doc = Document(docx_path)
    all_text_list = []

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    all_text_list.append(text)

    divided_data = [all_text_list[i:i + items_per_subarray] for i in range(0, len(all_text_list), items_per_subarray)]

    return divided_data

# Вставка данных в шаблон
def create_draft(data, output_docx_path, namepred, userscommission, userbd):
    doc = DocxTemplate('C:/Users/Acer/Desktop/VKR_finals/shablon.docx')
    
    combined_doc = Document()
    
    for item in data:
        date_obj = datetime.combine(item[0], datetime.min.time())
        day_name = date_obj.strftime('%d')
        month_name = date_obj.strftime('%B')
        year_name = date_obj.strftime('%Y')
        date_obj_end = datetime.combine(item[9], datetime.min.time())
        day_name_end = date_obj_end.strftime('%d')
        month_name_end = date_obj_end.strftime('%B')
        year_name_end = date_obj_end.strftime('%Y')
        context = {
            'day': day_name,
            'mouth': month_name,
            'year': year_name,
            'napravlenie': item[1],
            'predgos': namepred,
            'studentA': item[3],
            'title': item[4],
            'nauchruc': item[5],
            'rang': item[6],
            'student': item[3],
            'kval': item[7],
            'spec': item[8],
            'dayend': day_name_end,
            'mouthend': month_name_end,
            'yearend': year_name_end
        }

        for i, usercommission in enumerate(userscommission, start=1):
            parts = usercommission['user_name'].split()
            last_name = parts[0]
            initials = ''.join([name[0] + '.' for name in parts[1:]])
            shortened_name = f"{last_name} {initials}"
            context['namegoskom' + str(i)] = shortened_name

        parts = namepred.split()
        last_name = parts[0]
        initials = ''.join([name[0] + '.' for name in parts[1:]])
        shortened_name = f"{last_name} {initials}"
        context['predgossokr'] = shortened_name

        with get_db_connection() as conn:
            cursor = conn.cursor()
            cursor.execute('select uc.user_name, pi2.text_issues from protection_issues pi2, users_commission uc where pi2.id_student = %s and pi2.id_user_commission = uc.id_user', (item[2],))
            questions = cursor.fetchall()
            questions_dict = {}
            for question in questions:
                name = question[0]
                question = question[1]
                if name in questions_dict:
                    questions_dict[name].append(question)
                else:
                    questions_dict[name] = [question]
            questions = [{"name": name, "question": questions_dict[name]} for name in questions_dict]
            for i in range(len(questions)):
                context['question' + str(i+1)] = questions[i]['name']+ ' ' + str(' '.join(questions[i]['question']))

            cursor.execute('select ee.name, ee.text, ge.name from assessments a, existing_estimates ee, graduate_estimates ge where a.id_student = %s and a.score_graduate_work = ee.id and a.score_graduate = ge.id_estimates', (item[2],))
            score = cursor.fetchone()
            if score is None:
                context['score'] = ''
                context['scorediplom'] = ''
            else:
                context['score'] = score[0]
                context['scorediplom'] = score[2]
                lines = score[1].split(',')
                for i in range(len(lines)):
                    if i < len(lines) - 1:
                        context['haracteransver' + str(i+1)] = lines[i] + ','
                    else:
                        context['haracteransver' + str(i+1)] = lines[i]
            
            context['ekzscore'] = "не предусмотрен учебным планом"
            context['nodata'] = "–"
            data_start = str(item[0]).split('-')
            context['data'] = data_start[2] + '.' + data_start[1] + '.' + data_start[0] 
            
            context['studentU'] = item[3] 
            
            cursor.execute('select name from adminbd where id_userbd = %s',(userbd,)) 
            namesekret = cursor.fetchone()
            parts = namesekret[0].split()
            last_name = parts[0]
            initials = ''.join([name[0] + '.' for name in parts[1:]])
            shortened_name = f"{last_name} {initials}"
            context['sekretgossokr'] = shortened_name
                     
            
        doc.render(context)

        for element in doc.element.body:
            combined_doc.element.body.append(element)
            
    combined_doc.save(output_docx_path)


#Вствка данных в шаблон заседание кафедры
def create_draft_ZK(data, pred, json, adminuser, pps, checkedItems, output_path):
    docx = DocxTemplate('C:/Users/Acer/Desktop/VKR_finals/shablonZK.docx')
    context = {}
    date_obj = datetime.strptime(data, '%Y-%m-%d').date()
    date_time_obj = datetime.combine(date_obj, datetime.min.time())
    day_name = date_time_obj.strftime('%d')
    month_name = date_time_obj.strftime('%B')
    year_name = date_time_obj.strftime('%Y')
    context['day'] = day_name
    context['month'] = month_name
    context['year'] = year_name
    with get_db_connection() as conn:
        cursor = conn.cursor()
        cursor.execute('SELECT a.name, d.name, r.name, d.full_name FROM adminbd a, role r, departament d WHERE id_userbd = %s and a.id_role_pps = r.id and a.departament = d.id', (adminuser,))
        execute_admin = cursor.fetchone()
        context['name_kaf'] = execute_admin[1]
        parts = execute_admin[0].split()
        last_name = parts[0]
        initials = ''.join([name[0] + '.' for name in parts[1:]])
        shortened_name_admin = f"{last_name} {initials}"
        context['admin'] = f'{execute_admin[2].lower()} {shortened_name_admin}'
        parts = pred.split()
        last_name = parts[0]
        initials = ''.join([name[0] + '.' for name in parts[1:]])
        shortened_name = f"{last_name} {initials}"
        cursor.execute('select r."name" from pps p, role r where p."name" = %s and p.id_role = r.id', (pred,))
        execute_pps_pred = cursor.fetchone()
        context['chairperson'] = f'{execute_pps_pred[0].lower()} {shortened_name}'
        context['checkedItems'] = checkedItems
        role = execute_admin[1]
        cursor.execute('select p."name" ,r."name" from pps p, role r, departament d where p.id_role = r.id and d."name" = %s', (role,))
        execute_pps = cursor.fetchall()
        pps = []
        for item in execute_pps:
            parts = item[0].split()
            last_name = parts[0]
            initials = ''.join([name[0] + '.' for name in parts[1:]])
            shortened_name_pps = f"{last_name} {initials}"
            pps.append(f'{item[1].lower()} {shortened_name_pps}, ')
        pps.append(f'{execute_admin[2].lower()} {shortened_name_admin}')
        context['pps'] = pps
        context['countPPS'] = len(pps)
        array_text = []
        for item in json:
            if item['title'] != 'Выступили':
                text_dict = {
                    'key': item['key'],
                    'title': item['title'].upper(),
                    'titledraft': item['titledraft'],
                    'text': item['text']
                }
            else:
                za = item['za']
                prot = item['prot']
                vozd = item['vozd']
                golosovText = f'в голосовании участвовало {len(pps) - 1} человек:'
                golosovstring = f'«ЗА» - {za}; «против» - {prot}; «воздержались» - {vozd};'
                text_dict = {
                    
                    'key': '',
                    'title': item['title'].upper(),
                    'titledraft': item['titledraft'],
                    'text': item['text'],
                    'golosovTitle': 'Результаты голосования: ',
                    'golosovText': golosovText,
                    'stringGolosov': golosovstring
                }
            array_text.append(text_dict)
        context['mainText'] = array_text
        parts = pred.split()
        last_name = parts[0]
        initials = ''.join([name[0] + '.' for name in parts[1:]])
        shortened_name = f"{initials} {last_name}"
        context['zav_kaf'] = shortened_name
        parts = execute_admin[0].split()
        last_name = parts[0]
        initials = ''.join([name[0] + '.' for name in parts[1:]])
        shortened_name_admin = f"{initials} {last_name}"
        context['uchen_secr'] = shortened_name_admin
        context['name_kaf_full'] = execute_admin[3]

    docx.render(context)
    docx.save(output_path)
